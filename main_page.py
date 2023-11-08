# importamos las librerias
import streamlit as st
import datetime
import openai 
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
from decouple import config
from dotenv import load_dotenv
from email.message import EmailMessage
import ssl
import smtplib



# Empezamos a definir el contenido de la pagina
st.title('''
         **Semillero de Ciencias Sociales Computacionales - Tin:robot_face::bee:**
         ''') 
st.subheader('**Linea de investigación - Territorios Inteligentes**', divider= 'gray')  
st.subheader('**Facultad de Ciencias Sociales y Humanas** \n ### **:green[Universidad de Antioquia]**', divider= 'gray')  

st.markdown(''' ## **Modelo: Speech to text** ''')
'''
Con ayuda de este modelo, podras transcribir tus audios solamente siguiendo estas instrucciones:\n
1. Asegurate que los audios estén en formato .mp3, .m4a o .ogg y no pesen más de 20 MG.
2. Arrastra o adjunta tu archivo en la caja receptora.
3. Espera mientras termina de procesar el audio.
4. Una vez termine el proceso, descarga tu archivo .dox el cual contendrá el texto formateado.

Este y otros proyectos en construcción se sostendrán a lo largo del tiempo gracias a las donaciones.
Si deseas y puedes contribuir a la causa, cualquier monto es bienvenido. Sin embargo, si no estás en
posición de hacerlo en este momento, no te preocupes; siempre habrá otra oportunidad.
Por ahora, disfruta de esta fantástica aplicación.:loudspeaker::technologist: 
'''

# Cargar el archivo de audio
archivo_audio = st.file_uploader('Arrastra o ingresa tu archivo .mp3, .ma4, .ogg', type=['.mp3','.m4a', '.ogg'])
nombre_archivo: str = ''
# Verificar si se ha cargado un archivo
if archivo_audio is not None:
    nombre_archivo = archivo_audio.name
    # Abrir un archivo en modo escritura binaria ('wb') para guardar el archivo de audio
    
    with open(nombre_archivo, 'wb') as new_file:
        # Leer los datos del archivo cargado y escribirlos en el nuevo archivo
        new_file.write(archivo_audio.read())

    st.success(f'Archivo de audio "{nombre_archivo}" ha sido guardado exitosamente.')
    # No olvides manejar los casos en los que no se cargue un archivo o haya algún error.
    st.success(f'Archivo de audio "{nombre_archivo}" pronto estará procesandose.')

# Procesamiento del audio con Whisper-1

openai.api_key = config('API_KEY')
result: str = ''
list_transcripciones: dict = []
fecha_hora_actual = datetime.datetime.now()
fecha_hora = f"{fecha_hora_actual.strftime('%Y-%m-%d__%H:%M:%S')}"

# Abre el archivo de audio
if nombre_archivo:
    with open(nombre_archivo, "rb") as audio_file:
        resultado = openai.Audio.transcribe("whisper-1",
                                        audio_file,
                                        encoding="utf-8",
                                        response_format="text")
    
    list_transcripciones.append({'nombre_archivo': nombre_archivo,
                                'texto': resultado.strip(),
                                'fecha': fecha_hora,
                                'numero_palabras': len(resultado.strip().split())})

    print(f'El archivo: {nombre_archivo} ha sido procesado\n', sep='-->')
    print(list_transcripciones)

    st.success(f'Archivo de audio "{nombre_archivo}" ha sido procesado.')
    st.markdown(''' ## **Texto:** ''')

texto = ''
if len(list_transcripciones) > 0:
    texto = list_transcripciones[0]['texto']

    st.write(f'''
    {texto}
    ''')

# Agrega un párrafo al documento
encabezado: str = '''----------------------------------------------------------------------------------------------------------------------
Esta transcripción es producto de los desarrollos del equipo de
la Linea de insvtigación: Territorios Inteligentes, que hace parte del
grupo de insvtigación: Redes y Actores Sociales (RAS) del departamente de
Sociología de la Facultad de Ciencias Sociales y Humanas de la Universidad
de Antioquia. Esto es solo un modelo de prueba y por ende esta sujeto a errores.
----------------------------------------------------------------------------------------------------------------------'''

metadata: str = '''- Nombre de la trascritpción: {nombre_archivo}
- Fecha y hora en la que se realizó la transcripción: {fecha}
- Numero de palabras transcritas: {numero_palabras}

- Texto:'''

template = '''"{texto}"\t\n
----------------------------------------------------------------------------------------------------------------------'''

# Defino el nombre del documento
if len(list_transcripciones) > 0:
    nombre = list_transcripciones[0]['nombre_archivo'][:-7]
    nombre_archivo_docx = f"transcripcion_nombre_archivo_{nombre}.docx"

# Crea un nuevo documento
doc = docx.Document()

#doc.add_paragraph(encabezado)

# Agregar un título
titulo = doc.add_paragraph(encabezado, style= 'Body Text')
titulo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Alineación justificada

for dato in list_transcripciones:
  print(template.format(**dato))

  # Agregar metedatos

  metadatos_ = doc.add_paragraph(metadata.format(**dato), style= 'Body Text')
  metadatos_.bold = True

  # Agregar un párrafo y configurar la alineación como justificada
  paragraph = doc.add_paragraph(template.format(**dato), style= 'Body Text')
  paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Alineación justificada

 # doc.add_paragraph(template.format(**dato), style= 'Body Text' )
try:
# Guarda el documento en un archivo .docx
    if nombre_archivo_docx :
        doc.save(nombre_archivo_docx)
        st.success(f'Ya se ha creado su archivo .docx. Ahora puede dercargarlo dando click en "Download docx"')

        # Agregamos el boon de descarga
        with open(nombre_archivo_docx, "rb") as file:
            btn = st.download_button(
                    label="Download docx",
                    data=file,
                    file_name=nombre_archivo_docx,
                )
            
    
            # Esto es para enviar el archivo a un correo
            imail_emisor = config('CORREO_PERSONAL')
            imail_contraseña = config('GOOGLE_KEY')
            imail_receptor = config('CORREO_U')
            asunto = 'Archivo_traducción'
            cuerpo = 'Se adjunta archivo con la traducción'

            em = EmailMessage()
            em['From'] = imail_emisor
            em['To'] = imail_receptor
            em['Subject'] = asunto
            em.set_content(cuerpo)

            with open(nombre_archivo_docx, "rb") as f:
                em.add_attachment(
                    f.read(),
                    filename=nombre_archivo_docx,
                    maintype="application",
                    subtype="docx"
                )

            contexto = ssl.create_default_context()

            with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=contexto) as smtp:
                smtp.login(imail_emisor, imail_contraseña)
                smtp.sendmail(imail_emisor, imail_receptor, em.as_string()) 
                smtp.quit()

except:
    print('error')


st.markdown("### Donaciones:")

col1, col2, col3 = st.columns(3)

with col1:
    #st.subheader(':green[Invitame a un café]')
    boton_bmc = st.button('Link - buy me a coffe', use_container_width = True)
    if boton_bmc:
        st.markdown('https://www.buymeacoffee.com/cristianmoz') 
    bmec = 'complementos/bmc2.png'
    st.image(bmec)

with col2:
    #st.subheader(':green[Paypall]')
    boton_pp = st.button('link - Paypal', use_container_width = True)
    if boton_pp:
        st.markdown('https://paypal.me/CristianMontoya158?country.x=CO&locale.x=es_XC') 
    paypal = 'complementos/paypal1.png'
    st.image(paypal)

with col3:
    boto_bcol = st.button('QR - Bancolombia', use_container_width = True)
    bancolomia = 'complementos/bancolombia.png'
    st.image(bancolomia)

st.markdown(''' ## **Comentarios:** ''')
txt = st.text_area("Dejanos un comentario --- Presiona 'Ctrl + Enter' para enviar.")


# Si alguien deja un comentario, este se enviara al correo. Primero debe configurarlo
if txt:
    st.write(f'Comentario recibido.')

    print(txt)

    imail_emisor = config('CORREO_PERSONAL')
    imail_contraseña = config('GOOGLE_KEY')
    imail_receptor = config('CORREO_U')
    asunto = 'Comentario de un usuario de la aplicacion "audio_a_texto" de streamlit'
    cuerpo = txt

    em = EmailMessage()
    em['From'] = imail_emisor
    em['To'] = imail_receptor
    em['Subject'] = asunto
    em.set_content(cuerpo)

    contexto = ssl.create_default_context()

    with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=contexto) as smtp:
        smtp.login(imail_emisor, imail_contraseña)
        smtp.sendmail(imail_emisor, imail_receptor, em.as_string())   
    




